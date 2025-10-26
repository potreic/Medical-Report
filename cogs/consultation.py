import os
import re
import json
import aiohttp
import discord
from discord.ext import commands
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config.settings import TEMP_DIR, REPORT_DIR
from core.speech_to_text import transcribe_audio
from core.langchain_pipeline import generate_medical_report

# NOTE: Removed SESSION_TTL_MINUTES – patient info is one-time use per voice file

# ---------- style helpers ----------

def _set_page_margins(doc: Document, cm=2.0):
    for section in doc.sections:
        section.top_margin = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin = Cm(cm)
        section.right_margin = Cm(cm)


def _set_global_font(doc: Document, name="Times New Roman", size=12):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    if style.paragraph_format.line_spacing is None:
        style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)


def _set_cell_borders(cell, top="single", left="single", bottom="single", right="single", sz=12, color="000000"):
    """Set border untuk sebuah cell tabel."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    def _edge(tag, val):
        el = tcBorders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcBorders.append(el)
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)

    if top: _edge("top", top)
    if left: _edge("left", left)
    if bottom: _edge("bottom", bottom)
    if right: _edge("right", right)


def _shade_cell(cell, fill="EDEDED"):
    """Shading halus pada cell header (hex tanpa #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_box(doc: Document, title: str, lines: int = 6, page_width_cm: float = 16.0):
    """
   Create a box with a bold title at the top and lines for writing.
It is made using a single-column table: the first row is the header (with shading), and the remaining rows have bottom borders (for writing lines).
    """
    table = doc.add_table(rows=1 + lines, cols=1)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(page_width_cm)

    # Header
    hdr = table.cell(0, 0)
    _set_cell_borders(hdr, top="single", left="single", bottom="single", right="single", sz=16)
    _shade_cell(hdr)
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    # Garis isi
    for i in range(1, 1 + lines):
        c = table.cell(i, 0)
        _set_cell_borders(
            c,
            top="nil",
            left="single",
            bottom="single",
            right="single",
            sz=8
        )
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].add_run(" ")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _add_report_section_table(doc: Document, title: str, lines: list[str],
                              min_rows: int = 6, page_width_cm: float = 16.0):
    """A section box containing a header and rows with bottom borders."""
    total_rows = max(min_rows, max(1, len(lines)))
    table = doc.add_table(rows=1 + total_rows, cols=1)
    table.autofit = False
    for row in table.rows:
        row.cells[0].width = Cm(page_width_cm)

    # Header
    hdr = table.cell(0, 0)
    _set_cell_borders(hdr, top="single", left="single", bottom="single", right="single", sz=16)
    _shade_cell(hdr)
    p = hdr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    # Rows (garis bawah)
    for i in range(total_rows):
        c = table.cell(1 + i, 0)
        _set_cell_borders(c, top="nil", left="single", bottom="single", right="single", sz=8)
        para = c.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        if i < len(lines) and lines[i]:
            txt = re.sub(r"^[\-•]\s+", "", lines[i]).strip()
            para.add_run(txt)
        else:
            para.add_run(" ")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# ---------------------- template generator ----------------------


def save_consultation_template_with_boxes(path: str):
    """Create a .docx template file with outlines/boxes for each section."""
    doc = Document()
    _set_page_margins(doc, cm=2.0)
    _set_global_font(doc, size=12)

    _add_title(doc, "Medical Consultation Report")
    doc.add_paragraph()

    sections = [
        ("Symptoms", 6),
        ("Diagnosis", 4),
        ("Prescription / Treatment Plan", 8),
        ("Doctor's Notes", 6),
        ("Assessment", 4),
        ("Plan", 5),
        ("Red Flags", 4),
        ("Disclaimer", 3),
    ]

    for name, nlines in sections:
        _add_section_box(doc, name, lines=nlines)

    p = doc.add_paragraph()
    r = p.add_run(
        "Note: This document is generated from a clinician–patient consultation. "
        "It is intended for clinical use only and may contain transcription or inference errors."
    )
    r.italic = True
    r.font.size = Pt(10)

    doc.save(path)

# ---------------------- bot cog ----------------------


class Consultation(commands.Cog):
    """
Flow:
1) The user sends the patient's name and ID (using !patient ... or a formatted text message).
2)The user sends one audio file (mp3/wav/m4a/ogg) — the patient info is used once and then deleted.
3)The bot: downloads → transcribes → generates the report (.docx) → sends the file.
    """
    def __init__(self, bot):
        self.bot = bot
        self.audio_dir = TEMP_DIR
        self.transcript_dir = "data/transcripts"
        self.report_dir = REPORT_DIR
        os.makedirs(self.audio_dir, exist_ok=True)
        os.makedirs(self.transcript_dir, exist_ok=True)
        os.makedirs(self.report_dir, exist_ok=True)

        # In-memory, one-time use: { user_id: {"name": str, "id": str} }
        self.session_patients = {}

    # ---------------------- Helpers ----------------------

    @staticmethod
    def _parse_patient_info(text: str):
        """
     Extract the Name and ID from the message text.
    Supports multiple formats (case-insensitive).
        """
        if not text:
            return None, None
        t = text.strip()

        m1 = re.search(r'(?:name|nama)\s*=\s*"([^"]+)"', t, flags=re.IGNORECASE)
        m2 = re.search(r'(?:name|nama)\s*[:=]\s*"?([^\n";,]+?)"?\s*(?:$|[;,])', t, flags=re.IGNORECASE)
        name = None
        for m in (m1, m2):
            if m and m.group(1).strip():
                name = m.group(1).strip()
                break

        mid = re.search(r'(?:patient\s*id|id\s*pasien|id)\s*[:=]\s*([A-Za-z0-9\-_]+)', t, flags=re.IGNORECASE)
        pid = mid.group(1).strip() if mid else None

        return name, pid

    def _set_session_patient(self, user_id: int, name: str, pid: str):
        # No TTL; will be cleared after one successful audio processing
        self.session_patients[user_id] = {
            "name": name.strip(),
            "id": pid.strip(),
        }

    def _pop_session_patient(self, user_id: int):
        return self.session_patients.pop(user_id, None)

    def _get_session_patient(self, user_id: int):
        return self.session_patients.get(user_id)

    # ---------------------- Set patient via command (opsional) ----------------------

    @commands.command(name="patient")
    async def cmd_set_patient(self, ctx, *, args: str):
        name, pid = self._parse_patient_info(args)
        if not name or not pid:
            m = re.match(r'"\s*([^"]+)\s*"\s+([A-Za-z0-9\-_]+)', args)
            if m:
                name, pid = m.group(1), m.group(2)

        if not name or not pid:
            return await ctx.send(
                "Incorrect format.\nContoh:\n"
                "`!patient \"Kusuma\" P-00123`\n"
                "atau\n`!patient name=\"Kusuma\" id=P-00123`"
            )

        self._set_session_patient(ctx.author.id, name, pid)
        await ctx.send(
            f"Patient set for **{ctx.author.display_name}** (one-time use)\n"
            f"- Name: **{name}**\n- ID: **{pid}**\n"
            f"Send **one** file audio (mp3/wav/m4a/ogg);"
        )

    # ---------------------- Listener ----------------------

    @commands.Cog.listener()
    async def on_message(self, message: discord.Message):
        if message.author.bot:
            return


        if not message.attachments and (message.content and message.content.strip()):
            name, pid = self._parse_patient_info(message.content)
            if name and pid:
                self._set_session_patient(message.author.id, name, pid)
                return await message.channel.send(
                    "Patient info saved (one-time use).\n"
                    f"- Name: **{name}**\n- ID: **{pid}**\n"
                    "Send **one** file audio (mp3/wav/m4a/ogg)."
                )

      
        if not message.attachments:
            return

        audio_attachments = [
            a for a in message.attachments
            if a.filename.lower().endswith((".mp3", ".wav", ".m4a", ".ogg"))
        ]
        if not audio_attachments:
            return

        patient = self._get_session_patient(message.author.id)
        if not patient:
            return await message.channel.send(
                "**Patient info is required before sending audio.**\n"
                "send first Name and ID Patient (text message), contoh:\n"
                "• `Nama: \"Athaya Kusuma\"; ID: P-00123`\n"
                "• `name=\"Athaya Kusuma\" id=P-00123`\n"
                "or using command: `!patient \" Kusuma\" P-00123`"
            )

       
        if len(audio_attachments) > 1:
            await message.channel.send("Multiple audio files detected. Only the **first** file will be processed according to the one-time policy.")
        attachment = audio_attachments[0]

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_name = f"{message.author.name}_{timestamp}_{attachment.filename}"
        audio_path = os.path.join(self.audio_dir, safe_name)

        async with aiohttp.ClientSession() as session:
            async with session.get(attachment.url) as resp:
                if resp.status == 200:
                    with open(audio_path, "wb") as f:
                        f.write(await resp.read())

        await message.channel.send(
            f"Audio received: `{attachment.filename}`\n"
            f"Patient: **{patient['name']}**  |  ID: **{patient['id']}**\n"
            f"Transcribing..."
        )

        # Transcribe
        try:
            transcript_text = await self.bot.loop.run_in_executor(
                None, transcribe_audio, audio_path
            )
        except Exception as e:
            return await message.channel.send(f"Transcription failed: {e}")

        # Simpan transcript
        base = os.path.splitext(safe_name)[0]
        transcript_path = os.path.join(self.transcript_dir, f"{base}.txt")
        try:
            with open(transcript_path, "w", encoding="utf-8") as f:
                f.write(transcript_text)
        except Exception as e:
            return await message.channel.send(f"Saving transcript failed: {e}")

        await message.channel.send("Generating structured medical report...")

        # Generate medical report
        try:
            report_text = await self.bot.loop.run_in_executor(
                None, generate_medical_report, transcript_text
            )
        except Exception as e:
            return await message.channel.send(f"Report generation failed: {e}")

    
        report_docx_path = os.path.join(self.report_dir, f"{base}.docx")
        try:
            save_tidy_docx(
                report_text=report_text,
                report_path=report_docx_path,
                author_name=message.author.display_name,
                patient_name=patient["name"],
                patient_id=patient["id"],
            )
        except Exception as e:
            return await message.channel.send(f"Saving DOCX failed: {e}")

        # One-time use: clear immediately after successful generation
        self._pop_session_patient(message.author.id)

        await message.channel.send(
            "Report generated successfully! (patient info cleared)",
            file=discord.File(report_docx_path)
        )

# ---------------------- DOCX Renderer ----------------------


def _looks_like_patient_info(line: str) -> bool:
    """A filter to prevent model content containing patient identity information from being rewritten in the body."""
    patterns = [
        r"^\s*(patient\s*name|nama\s*pasien)\s*[:=]",
        r"^\s*(patient\s*id|id\s*pasien|mrn|rekam\s*medis)\s*[:=]",
        r"^\s*(dob|tanggal\s*lahir|age|umur)\s*[:=]",
        r"^\s*(date\s*of\s*consultation|tanggal\s*konsultasi)\s*[:=]",
        r"^\s*(alamat|address|phone|telepon)\s*[:=]",
    ]
    for p in patterns:
        if re.search(p, line, flags=re.IGNORECASE):
            return True
    return False


def _add_divider(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("\u2500" * 60)  # garis tipis
    run.font.size = Pt(8)

# ---- robust section header matching ----
SECTION_ALIASES = {
    "Symptoms": r"symptoms?",
    "Diagnosis": r"(diagnosis|dx)",
    "Prescription / Treatment Plan": r"(prescription|rx|treatment\s*/?\s*plan|management\s*plan)",
    "Doctor's Notes": r"(doctor'?s?\s*notes?|clinical\s*notes?)",
    "Assessment": r"(assessment|impression)",
    "Plan": r"plan(?:\s*of\s*care)?",
    "Red Flags": r"(red\s*flags?|warning\s*signs?)",
    "Disclaimer": r"(disclaimer|note)",
}


def _canonical_section(line: str, sections: list[str]) -> str | None:
    """ChatGPT said:

Return the canonical section name if the line matches the header."""
    if not line:
        return None
    norm = re.sub(r"\s+", " ", line.strip()).rstrip(":").lower()
    norm = re.sub(r"\s*/\s*", "/", norm)  # standardize ' / '

    # Exact match vs canonical header (with normalized '/')
    for s in sections:
        s_norm = re.sub(r"\s*/\s*", "/", s.lower())
        if norm == s_norm:
            return s

    # Alias match
    for canonical, pat in SECTION_ALIASES.items():
        if re.fullmatch(pat, norm, flags=re.IGNORECASE):
            return canonical
    return None


def save_tidy_docx(
    report_text: str,
    report_path: str,
    author_name: str = "Unknown",
    patient_name: str = "Unknown",
    patient_id: str = "Unknown",
):
    """
  Generate a clean, professional medical report in Times New Roman (.docx).
 -Inserts the patient's Name and ID only at the top (without duplication in the body).
 -Detects the consultation date from the filename (YYYYMMDD) if available.
 -BODY: each section is rendered as a bordered box like a form.
    """
    # Extract date from filename if possible (…_YYYYMMDD_HHMMSS_…)
    m = re.search(r"(\d{8})_\d{6}", report_path)
    if m:
        raw = m.group(1)
        try:
            consultation_date = datetime.strptime(raw, "%Y%m%d").strftime("%d %B %Y")
        except ValueError:
            consultation_date = "[Date not available]"
    else:
        consultation_date = "[Date not available]"

    doc = Document()
    _set_page_margins(doc, cm=2.0)
    _set_global_font(doc, size=12)

    # Title
    _add_title(doc, "Medical Consultation Report")
    doc.add_paragraph()

    # Header info (patient + date)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs.append(header.add_run(f"Patient Name: {patient_name}\n"))
    header.runs.append(header.add_run(f"Patient ID: {patient_id}\n"))
    header.runs.append(header.add_run(f"Date of Consultation: {consultation_date}\n"))
    header.runs.append(header.add_run(f"Generated by: {author_name}\n"))

    _add_divider(doc)

    # Definisi section yang dikenali
    sections = [
        "Symptoms",
        "Diagnosis",
        "Prescription / Treatment Plan",
        "Doctor's Notes",
        "Assessment",
        "Plan",
        "Red Flags",
        "Disclaimer",
    ]

    buckets = {s: [] for s in sections}
    current = None
    for raw in [ln.strip() for ln in report_text.splitlines() if ln.strip()]:
        maybe = _canonical_section(raw, sections)
        if maybe:
            current = maybe
            continue
        if current is None:
            current = "Doctor's Notes"
        if _looks_like_patient_info(raw):
            continue
        line = re.sub(r"^[\-•]\s+", "- ", raw)
        buckets[current].append(line)

    # Render tiap section sebagai kotak bergaris
    MIN_ROWS = {
        "Symptoms": 6,
        "Diagnosis": 4,
        "Prescription / Treatment Plan": 8,
        "Doctor's Notes": 6,
        "Assessment": 4,
        "Plan": 5,
        "Red Flags": 4,
        "Disclaimer": 3,
    }

    for s in sections:
        content = buckets.get(s) or ["- None reported."]
        _add_report_section_table(doc, s, content, min_rows=MIN_ROWS.get(s, 5))

    doc.save(report_path)


async def setup(bot):
    await bot.add_cog(Consultation(bot))
